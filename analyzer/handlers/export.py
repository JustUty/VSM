import io
import re
from datetime import datetime

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font

from analyzer.handlers.human_readable import get_human_message_templates


def clean_text(text):
    """Очищает текст от недопустимых символов для XML"""
    if text is None:
        return ''
    text = str(text)
    if pd.isna(text) or text == 'NaT' or text == 'nan':
        return ''
    text = text.replace('\x00', '')
    text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    if len(text) > 500:
        text = text[:497] + '...'
    return text


def format_datetime(dt):
    """Безопасное форматирование даты/времени"""
    if dt is None:
        return ''
    if pd.isna(dt):
        return ''
    if isinstance(dt, (datetime, pd.Timestamp)):
        try:
            return dt.strftime("%d.%m.%Y %H:%M:%S")
        except Exception:
            return str(dt)
    return str(dt)

def format_car_number(carnumber):
    """Преобразует номер вагона в формат +100, +200, +300 и т.д."""
    car = clean_text(carnumber)

    if not car:
        return ''

    last_two = car[-2:]

    if last_two.isdigit():
        car_index = int(last_two)
        if car_index > 0:
            return f'+{car_index * 100}'

    return car


def add_brackets_to_ds_code(text, code):
    """Добавляет квадратные скобки к коду ДС в тексте сообщения."""
    text = clean_text(text)
    code = clean_text(code)

    if not text or not code:
        return text

    text = re.sub(
        rf'\bДС\s+{re.escape(code)}\b',
        f'ДС [{code}]',
        text
    )

    if f'[{code}]' not in text and code in text:
        text = text.replace(code, f'[{code}]', 1)

    return text


def get_protocol_message_text(row):
    """Получает текст сообщения для человекочитаемого протокола."""

    code = clean_text(row.get('messagecode', ''))
    event_type = clean_text(row.get('event_type', ''))
    message_text = clean_text(row.get('message_text', ''))

    templates = get_human_message_templates(code)

    # Выбираем нужный шаблон
    if event_type == 'activation':
        text = templates.get('kurztext_3') or message_text

    elif event_type == 'deactivation':
        text = templates.get('kurztext_4') or message_text

    else:
        text = message_text or templates.get('kurztext_2', '')

    text = clean_text(text)

    # Если есть "ДС 44051" → делаем "ДС [44051]"
    text = re.sub(
        rf'\bДС\s+{re.escape(code)}\b',
        f'ДС [{code}]',
        text
    )

    # Если код уже есть в квадратных скобках — ничего не делаем
    if f'[{code}]' in text:
        return text

    # Если код вообще отсутствует — добавляем в начало
    return f'[{code}] {text}'


def build_human_readable_entry(row):
    """Формирует человекочитаемый текстовый блок для одной записи"""
    train_id = clean_text(row.get('train_id', ''))
    carnumber = format_car_number(row.get('carnumber', ''))
    timestamp = row.get('timestamp', None)

    timestamp_str = format_datetime(timestamp)
    message_text = get_protocol_message_text(row)

    lines = []

    header_parts = []

    if train_id:
        header_parts.append(f'поезд {train_id}')

    if carnumber:
        header_parts.append(f'вагон {carnumber}')

    if header_parts:
        lines.append(', '.join(header_parts) + '.')

    if timestamp_str and message_text:
        lines.append(f'{timestamp_str} {message_text}.')
    elif timestamp_str:
        lines.append(f'{timestamp_str} Зафиксировано диагностическое сообщение.')

    return '\n'.join(lines)


def build_human_readable_protocol_text(timeline_df, train_human_name, dt_from, dt_to):
    """Формирует полный человекочитаемый протокол"""
    lines = [
        "Эксплуатационный протокол",
        "",
        f"Поезд: {train_human_name}",
        f"Период: с {format_datetime(dt_from)} по {format_datetime(dt_to)}",
        f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        "",
    ]

    if timeline_df.empty:
        lines.append("За указанный период диагностические события не обнаружены.")
        return '\n'.join(lines)

    for _, row in timeline_df.iterrows():
        entry_text = build_human_readable_entry(row)
        lines.append(entry_text)
        lines.append("")

    return '\n'.join(lines).strip()


def export_text_to_docx(protocol_text, file_title="Эксплуатационный протокол"):
    """Экспортирует произвольный отредактированный текст протокола в DOCX"""
    try:
        doc = Document()

        lines = str(protocol_text).splitlines()

        if lines:
            first_line = lines[0].strip()
            if first_line:
                title = doc.add_heading(first_line, 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                lines = lines[1:]

        for line in lines:
            if line.strip():
                doc.add_paragraph(clean_text(line))
            else:
                doc.add_paragraph("")

        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
    except Exception as e:
        print(f"EDITED DOCX export error: {e}")
        return None


def get_column_names_map():
    """Возвращает маппинг колонок для отображения"""
    return {
        'train_id': 'Номер поезда',
        'carnumber': 'Вагон',
        'messagecode': 'Код ДС',
        'event_type': 'Тип события',
        'timestamp': 'Время события',
        'message_text': 'Описание',
        'duration_str': 'Продолжительность',
        'parsingtime': 'Время парсинга'
    }


def prepare_row_for_export(row, col):
    """Подготавливает значение ячейки для экспорта"""
    value = row.get(col, '')

    if col == 'event_type':
        if value == 'activation':
            return 'Активация'
        elif value == 'deactivation':
            return 'Деактивация'
        elif value == 'still_active_marker':
            return 'Активно до сих пор'
        else:
            return '—'
    elif col == 'timestamp':
        return format_datetime(value)
    else:
        return clean_text(value)


def export_to_docx(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    """Экспорт хронологии в DOCX (только выбранные колонки)"""
    try:
        doc = Document()

        title = doc.add_heading('Эксплуатационный протокол', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Поезд: {train_human_name}')
        doc.add_paragraph(f'Период: с {format_datetime(dt_from)} по {format_datetime(dt_to)}')
        doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('')

        # Если не выбраны колонки — используем стандартные
        if not selected_columns:
            selected_columns = ['train_id', 'carnumber', 'messagecode', 'event_type', 'timestamp', 'message_text']

        column_names = get_column_names_map()

        # Заголовки
        headers = [column_names.get(col, col) for col in selected_columns if col in timeline_df.columns]
        selected_cols = [col for col in selected_columns if col in timeline_df.columns]

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Данные
        for _, row in timeline_df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(selected_cols):
                cells[i].text = str(prepare_row_for_export(row, col))

        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
    except Exception as e:
        print(f"DOCX export error: {e}")
        return None


def export_human_readable_docx(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    """Экспорт хронологии в DOCX (человекочитаемый формат)"""
    try:
        doc = Document()

        title = doc.add_heading('Эксплуатационный протокол', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Поезд: {train_human_name}')
        doc.add_paragraph(f'Период: с {format_datetime(dt_from)} по {format_datetime(dt_to)}')
        doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        doc.add_paragraph('')

        if timeline_df.empty:
            doc.add_paragraph('За указанный период диагностические события не обнаружены.')
        else:
            for _, row in timeline_df.iterrows():
                entry_text = build_human_readable_entry(row)

                for line in entry_text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(clean_text(line))

                doc.add_paragraph('')

        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes
    except Exception as e:
        print(f"HUMAN DOCX export error: {e}")
        return None


def export_to_xlsx(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    """Экспорт хронологии в XLSX (только выбранные колонки)"""
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Эксплуатационный протокол"

        # Информационная шапка
        ws['A1'] = f'Поезд: {train_human_name}'
        ws['A2'] = f'Период: с {format_datetime(dt_from)} по {format_datetime(dt_to)}'
        ws['A3'] = f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}'
        ws['A5'] = ''

        # Если не выбраны колонки — используем стандартные
        if not selected_columns:
            selected_columns = ['train_id', 'carnumber', 'messagecode', 'event_type', 'timestamp', 'message_text']

        column_names = get_column_names_map()

        # Заголовки
        headers = [column_names.get(col, col) for col in selected_columns if col in timeline_df.columns]
        selected_cols = [col for col in selected_columns if col in timeline_df.columns]

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=6, column=col, value=header)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')

        # Данные
        for row_idx, (_, row) in enumerate(timeline_df.iterrows(), 7):
            for col_idx, col in enumerate(selected_cols, 1):
                ws.cell(row=row_idx, column=col_idx, value=prepare_row_for_export(row, col))

        # Автоматическая ширина колонок
        for col in range(1, len(headers) + 1):
            max_length = 0
            column_letter = ws.cell(row=6, column=col).column_letter
            for row in range(6, ws.max_row + 1):
                cell_value = ws.cell(row=row, column=col).value
                if cell_value:
                    max_length = max(max_length, len(str(cell_value)))
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width

        xlsx_bytes = io.BytesIO()
        wb.save(xlsx_bytes)
        xlsx_bytes.seek(0)
        return xlsx_bytes
    except Exception as e:
        print(f"XLSX export error: {e}")
        return None


def export_to_csv(timeline_df, train_human_name, dt_from, dt_to, selected_columns=None):
    """Экспорт хронологии в CSV (только выбранные колонки)"""
    try:
        # Если не выбраны колонки — используем стандартные
        if not selected_columns:
            selected_columns = ['train_id', 'carnumber', 'messagecode', 'event_type', 'timestamp', 'message_text']

        # Выбираем только нужные колонки
        available_columns = [col for col in selected_columns if col in timeline_df.columns]
        df_clean = timeline_df[available_columns].copy()

        # Преобразуем event_type в читаемый вид
        if 'event_type' in df_clean.columns:
            event_type_map = {
                'activation': 'Активация',
                'deactivation': 'Деактивация',
                'still_active_marker': 'Активно до сих пор'
            }
            df_clean['event_type'] = df_clean['event_type'].map(event_type_map).fillna(df_clean['event_type'])

        # Форматируем timestamp
        if 'timestamp' in df_clean.columns:
            df_clean['timestamp'] = df_clean['timestamp'].apply(
                lambda x: format_datetime(x) if x is not None else ''
            )

        # Очищаем текстовые колонки
        for col in df_clean.columns:
            if df_clean[col].dtype == 'object':
                df_clean[col] = df_clean[col].apply(
                    lambda x: clean_text(x) if x is not None else ''
                )

        csv_data = io.StringIO()
        df_clean.to_csv(csv_data, index=False, encoding='utf-8-sig')
        csv_bytes = io.BytesIO(csv_data.getvalue().encode('utf-8-sig'))
        return csv_bytes
    except Exception as e:
        print(f"CSV export error: {e}")
        return None